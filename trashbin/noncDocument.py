# -*- coding: utf-8 -*-
"""
Created on Mon Dec  3 09:46:06 2018

@author: alchenerd (alchenerd@gmail.com)
"""

from docx import Document

class NoncDocument():
    def __init__(self):
        self.templates = {
                "test": self._build_test}

    def write_template(self, doc:Document, mode:str):
        return self.templates.get(mode)(doc)

    def _build_test(self, d:Document):
        d.add_paragraph('Test')
        return d

def main():
    d = Document()
    noncd = NoncDocument()
    noncd.write_template(d, "test")
    print(type(d.styles))
    d.save("nonc.docx")

if __name__ == '__main__':
    main()