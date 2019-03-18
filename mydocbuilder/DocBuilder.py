# -*- coding: utf-8 -*-
"""
@author: alchenerd (alchenerd@gmail.com)
"""

import datetime
import os
import sqlite3
import sys
from copy import copy, deepcopy

import comtypes
import comtypes.client
import deprecation
import xlwt
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt
from PyQt5.QtCore import pyqtSignal, QObject

from myconnect import connect


if __name__ == '__main__':  # at mydocbuilder
    sys.path.append('../')


wdFormatPDF = 17  # magic constant but I'm too lazy to change case


class DocBuilder(QObject):
    """DocBuilder is a customized docx creator exclusively for hvhnonc."""

    # emits max, current, msg
    status_update = pyqtSignal(int, int, str)

    def __init__(self, type_: str = 'default', **kwargs):
        super(QObject, self).__init__()
        self.actions = {
            'default': self.hello_docx,
            'register_list': self.create_register_list,
            'unregister_list': self.create_unregister_list,
            'monthly_report': self.create_monthly_report,
            'full_report': self.create_full_report}
        if self.actions.get(type_, None):
            self.type_ = type_
        else:
            self.type_ = 'default'
            self.kwargs.clear()
        self.kwargs = kwargs.get('kwargs')

    def set_type(self, type_: str = 'default'):
        """Sets the type of which kind of document is going to be built."""
        self.type_ = type_

    def set_kwargs(self, **kwargs):
        """Sets the form info into the doc builder."""
        self.kwargs = deepcopy(kwargs)

    def construct(self):
        """Constructs a docx and save it."""
        # see self.actions for individual construct functions
        self.actions[self.type_]()

    def hello_docx(self):
        """Makes a dummy hello docx document."""
        self.status_update.emit(1, 1, "You shouldn't be seeing this, hmm.")
        document = Document()
        document.add_heading('Hello .docx!', 0)
        p = document.add_paragraph('This is my test paragraph!')
        records = ((2, 9, 4), (7, 5, 3), (6, 1, 8))
        table = document.add_table(rows=0, cols=3)
        for x, y, z in records:
            rowCells = table.add_row().cells
            rowCells[0].text = str(x)
            rowCells[1].text = str(y)
            rowCells[2].text = str(z)
        document.save('result.docx')

    def setMyFont(self, doc):
        """Set normal font of doc as my font."""
        doc.styles['Normal'].font.name = u'標楷體'
        doc.styles['Normal'].font.size = Pt(12)
        doc.styles['Normal']._element.rPr.rFonts.set(
            qn('w:eastAsia'), u'標楷體')

    def docx_to_pdf(self, in_file, out_file):
        comtypes.CoInitialize()
        word = comtypes.client.CreateObject('Word.Application')
        docx = word.Documents.Open(in_file)
        docx.SaveAs(out_file, FileFormat=wdFormatPDF)
        docx.Close()
        word.Quit()

    def create_register_list(self):
        """Creates a register list, saves data as excel, docx, and pdf."""

        def fetch_from_database(d):
            """Returns a list of sqlite3.Row as data"""
            con, cur = connect._get_connection(useSQL3Row=True)
            sqlstr = ('select {columns} from {table} where {conditions}')
            replacements = {}
            # columns: object_ID, serial_ID, name, spec, unit, amount, price,
            #          acquire_date, keep_year, keep_department, place, keeper
            replacements['columns'] = ', '.join((
                'ID', 'object_ID', 'serial_ID', 'name', 'spec', 'unit',
                'amount', 'price', 'acquire_date', 'keep_year',
                'keep_department', 'place', 'keeper'))
            # table: hvhnonc_in
            replacements['table'] = 'hvhnonc_in'
            # conditions: determined by d
            replacements['conditions'] = ''
            params = []
            for k, v in d.items():
                if 'date' in k:
                    # date string tuple
                    replacements['conditions'] += \
                        '({} between ? and ?) and '.format(k)
                    params.extend(v)
                else:
                    replacements['conditions'] += \
                        ('{0} like ? and '.format(k))
                    params.append('%' + v + '%')
            replacements['conditions'] += '1'
            # fill in the blanks
            sqlstr = sqlstr.format(**replacements)
            cur.execute(sqlstr, params)
            data = cur.fetchall()
            con.close()
            return data

        def parse_for_document(rows):
            """Parse sqlite3 rows to list of list for document table uses."""
            if len(rows) == 0:
                return [[]]
            result = []
            colTitle = ['物品編號', '物品名稱', '規格', '單位', '數量', '單價',
                        '總價', '取得日期', '使用年限', '存置地點',
                        '保管或使用單位', '保管或使用人']
            result.append(colTitle)
            # data
            for row in rows:
                # result row
                rrow = []
                # rrow[0]: objid + serial
                obj_ID = row['object_ID'].replace(' ', '')
                s = '-'.join((obj_ID, row['serial_ID']))
                rrow.append(s)
                # rrow[1:6]: name, spec, unit, amount, price
                rrow.append(str(row['name']))
                rrow.append(str(row['spec']))
                rrow.append(str(row['unit']))
                rrow.append(str(row['amount']))
                rrow.append(str(row['price']))
                # rrow[6]: total price
                rrow.append(str(row['amount'] * row['price']))
                # rrow[7]: acquire_date(EE/mm/dd)
                date = list(map(int, row['acquire_date'].split('-')))
                date[0] = date[0] - 1911
                date = list(map(lambda x: str(x).zfill(2), date))
                rrow.append('/'.join(date))
                # rrow[8:12]: keep_year, keep_department, place, keeper
                rrow.append(str(row['keep_year']))
                rrow.append(str(row['place']))
                rrow.append(str(row['keep_department']))
                rrow.append(str(row['keeper']))
                result.append(rrow)
            print(result)
            return result

        def write_to_excel(arr, fn):
            """Save 2d list to result.excel."""
            wb = xlwt.Workbook()
            ws = wb.add_sheet('result')
            for rc, row in enumerate(arr):
                for cc, column in enumerate(row):
                    try:
                        ws.write(r=rc, c=cc, label=column)
                    except:
                        # skip if encounter problems
                        continue
            wb.save(fn)

        def update_add_list_id(data):
            """Write add_list_ID to database.
            For now we use document creation date(today)."""
            today = datetime.date.today()
            today_str_zfill = (str(today.year), str(today.month).zfill(2),
                                str(today.day).zfill(2))
            string_today = '-'.join(today_str_zfill)
            con, cur = connect._get_connection()
            for row in data:
                id = row['ID']
                sqlstr = ('update hvhnonc_in set add_list_ID = ? where ID = ?')
                params = (string_today, id)
                cur.execute(sqlstr, params)
            con.commit()
            con.close()

        def construct_docx(data, status_update):
            """Open template, then modify according to rowCount."""
            doc = Document('./mydocbuilder/register_list_template.docx')
            # set font to 標楷體(16)
            self.setMyFont(doc)
            # fill in the header
            header = doc.sections[0].header
            replaceRow = header.tables[0].rows[0]
            # fill in department
            target_paragraph = replaceRow.cells[0].paragraphs[0]
            target_paragraph.text = \
                target_paragraph.text.format(**{'dept': '秘書室'})
            target_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # fill in the date
            target_paragraph = replaceRow.cells[1].paragraphs[0]
            today = datetime.date.today()
            s = [str(today.year - 1911), str(today.month).zfill(2),
                 str(today.day).zfill(2)]
            s = '中華民國{0}年{1}月{2}日'.format(*s)
            target_paragraph.text = s
            target_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for i, datum in enumerate(data[1:]):
                row = doc.tables[0].add_row()
                for cc, cell in enumerate(row.cells):
                    cell.paragraphs[0].text = datum[cc]
                self.status_update.emit(
                        6, 4,
                        'writing word file({}/{})...'.format(i, len(data) - 2))
            return doc

        self.status_update.emit(6, 0, 'initalizing...')  #0/6
        #print('create_register_list')
        # fetch data from database
        self.status_update.emit(6, 1, 'fetching...')  #1/6
        data = fetch_from_database(self.kwargs)
        # parse data for xls, docx
        self.status_update.emit(6, 2, 'parsing...')  #2/6
        data_parsed = parse_for_document(data)
        # write data and save to excel
        self.status_update.emit(6, 3,'writing excel file...') #3/6
        write_to_excel(data_parsed, 'result.xls')
        # write to docx template
        #status update 4/6 is in the function
        document = construct_docx(data_parsed, self.status_update)
        # save .docx
        document.save('result.docx')
        # convert to pdf and save (using current working directory)
        self.status_update.emit(6, 5,'converting to pdf...') #5/6
        cwd = os.getcwd()
        self.docx_to_pdf(cwd + '\\\\result.docx', cwd + '\\\\result.pdf')
        # update add_list_ID using sql Row data
        update_add_list_id(data)
        self.status_update.emit(6, 6,'Done!') #6/6

    def create_unregister_list(self):
        """Creates an unregister list, saves data as excel, docx, and pdf."""

        def fetch_from_database(d):
            """Returns a list of sqlite3.Row as data"""
            print(d)
            con, cur = connect._get_connection(useSQL3Row=True)
            con.set_trace_callback(print)
            sqlstr = (
                'select {columns} '
                'from {intable} as i '
                'inner join {outtable} as o '
                'on i.ID=o.in_ID '
                'and {conditions}')
            replacements = {}
            # columns: object_ID, serial_ID, name, spec, unit,
            #          unregister amount, date, keep year, used year, reason,
            #          remark
            replacements['columns'] = (
                'i.object_ID as object_ID, i.serial_ID as serial_ID, '
                'i.name as name, i.spec as spec, i.unit as unit, '
                'o.amount as amount, o.unregister_date as unregister_date, '
                'i.keep_year as keep_year, i.acquire_date as acquire_date, '
                'o.reason as reason, o.unregister_remark as remark')
            replacements['intable'] = 'hvhnonc_in'
            replacements['outtable'] = 'hvhnonc_out'
            # conditions: determined by d
            replacements['conditions'] = ''
            params = []
            for k, v in d.items():
                if 'date' in k:
                    if k == 'acquire_date':  # special case
                        # date string tuple
                        replacements['conditions'] += \
                            '(unregister_date between ? and ?) and '
                        params.extend(v)
                    else:
                        # date string tuple
                        replacements['conditions'] += \
                            '({} between ? and ?) and '.format(k)
                        params.extend(v)
                else:
                    replacements['conditions'] += \
                        ('{0} like ? and '.format(k))
                    params.append('%' + v + '%')
            replacements['conditions'] += '1'
            # fill in the blanks
            sqlstr = sqlstr.format(**replacements)
            print(sqlstr)
            cur.execute(sqlstr, params)
            data = cur.fetchall()
            con.close()
            for row in data:
                print(', '.join([str(row[k]) for k in row.keys()]))
            return data

        def parse_for_document(rows):
            """Parse sqlite3 rows to list of list for document table uses."""
            if len(rows) == 0:
                return [[]]
            result = []
            colTitle = ['物品編號', '物品名稱', '規格', '單位', '數量', '取得日期',
                        '使用年限', '已使用期間', '報廢原因', '審核意見', '備註']
            result.append(colTitle)
            # data
            for row in rows:
                # result row
                rrow = []
                # rrow[0]: objid + serial
                obj_ID = row['object_ID'].replace(' ', '')
                s = '-'.join((obj_ID, row['serial_ID']))
                rrow.append(s)
                # rrow[1:5]: name, spec, unit, amount
                rrow.append(str(row['name']))
                rrow.append(str(row['spec']))
                rrow.append(str(row['unit']))
                rrow.append(str(row['amount']))
                # rrow[5]: acquire_date(EE/mm/dd)
                date = list(map(int, row['acquire_date'].split('-')))
                date[0] = date[0] - 1911
                date = list(map(lambda x: str(x).zfill(2), date))
                rrow.append('/'.join(date))
                # rrow[6]: keep_year
                rrow.append(str(row['keep_year']))
                # rrow[7]: time used: '(y/m)'
                # used time(in months) calculate in acquire and retire date
                acqY, acqM, acqD = map(int, row['acquire_date'].split('-'))
                retY, retM, retD = map(int, row['unregister_date'].split('-'))
                hasRemain = int(retD - acqD > 0)
                detY = retY - acqY
                detM = retM - acqM + hasRemain
                if detM < 0:
                    detY -= 1
                    detM += 12
                delta = (str(detY), str(detM))
                delta = '(' + '/'.join(delta) + ')'
                rrow.append(delta)
                # rrow[8]: reason
                rrow.append(str(row['reason']))
                # rrow[9]: approval remarks, left blank
                rrow.append('')
                # rrow[10]: remark
                rrow.append(str(row['remark']))
                result.append(rrow)
            print(result)
            return result

        def write_to_excel(arr, fn):
            """Save 2d list to result.xls."""
            wb = xlwt.Workbook()
            ws = wb.add_sheet('result')
            for rc, row in enumerate(arr):
                for cc, column in enumerate(row):
                    try:
                        ws.write(r=rc, c=cc, label=column)
                    except:
                        # skip if encounter problems
                        continue
            wb.save(fn)

        def construct_docx(data, status_update):
            """Open template, then modify according to rowCount."""
            doc = Document('./mydocbuilder/unregister_list_template.docx')
            # set font to 標楷體(12)
            self.setMyFont(doc)
            # fill in the header
            header = doc.sections[0].header
            replaceRow = header.tables[0].rows[0]
            # fill in department
            target_paragraph = replaceRow.cells[0].paragraphs[0]
            target_paragraph.text = \
                target_paragraph.text.format(**{'dept': '秘書室'})
            target_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # fill in the date
            target_paragraph = replaceRow.cells[1].paragraphs[0]
            today = datetime.date.today()
            s = [str(today.year - 1911), str(today.month).zfill(2),
                 str(today.day).zfill(2)]
            s = '中華民國{0}年{1}月{2}日'.format(*s)
            target_paragraph.text = s
            target_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for i, datum in enumerate(data[1:]):
                row = doc.tables[0].add_row()
                for cc, cell in enumerate(row.cells):
                    cell.paragraphs[0].text = datum[cc]
                status_update.emit(
                        6, 4,
                        'writing word file({}/{})...'.format(i, len(data) - 2))
            return doc

        self.status_update.emit(6, 0, 'initalizing...')  #0/6
        #print('create_unregister_list')
        # fetch data from database
        self.status_update.emit(6, 1, 'fetching...')  #1/6
        data = fetch_from_database(self.kwargs)
        # parse data for xls, docx
        self.status_update.emit(6, 2, 'parsing...')  #2/6
        data = parse_for_document(data)
        # write data and save to excel
        self.status_update.emit(6, 3, 'writing excel file...')  #3/6
        write_to_excel(data, 'result.xls')
        # write to docx template
        # status_update 4/6 is done in the function
        document = construct_docx(data, self.status_update)
        # save .docx
        document.save('result.docx')
        # convert to pdf and save (using current working directory)
        self.status_update.emit(6, 5, 'converting to pdf...')  #5/6
        cwd = os.getcwd()
        self.docx_to_pdf(cwd + '\\\\result.docx', cwd + '\\\\result.pdf')
        self.status_update.emit(6, 6, 'Done!')  #6/6

    def create_monthly_report(self):
        """Creates an monthly report, saves data as excel, docx, and pdf."""

        def fetch_from_database(d):
            """fetch p1 and detail data from database.

            Args:
                d: The dictionary which is useful in forging sql conditions
            """
            yield fetch_p1(d)
            yield fetch_details(d)

        def get_year_month_int(d):
            try:
                _, magicDate = d.get('acquire_date')
                return(magicDate.year, magicDate.month)
            except (TypeError, ValueError, IndexError, NameError):
                today = datetime.datetime.today()
                return (today.year, today.month)

        def get_sql_conditions(d):
            # conditions: determined by d
            replacements = {'conditions': ''}
            params = []
            for k, v in d.items():
                # date string tuple
                if 'date' in k:
                    if k == 'acquire_date':  # ignore acquire date
                        pass
                    else:
                        replacements['conditions'] += \
                            '({} between ? and ?) and '.format(k)
                        params.extend(v)
                else:
                    replacements['conditions'] += \
                        ('{0} like ? and '.format(k))
                    params.append('%' + v + '%')
            replacements['conditions'] += '1'
            return replacements, params

        def fetch_p1(d):
            """The sql rows for the 1st page.

            Yields: income before month,
                    expense before month,
                    income in month,
                    expense in month
                    """
            # income before
            con, cur = connect._get_connection(useSQL3Row=True)
            sqlstr = (
                'select '
                'description as key, '
                'coalesce(sum(price * amount), 0) as value '
                'from hvhnonc_category '
                'left join hvhnonc_in '
                'on description = category '
                'and acquire_date < ? '
                'and {conditions} '
                'group by description;')
            replacements, params = get_sql_conditions(d)
            params.insert(0, datetime.date(
                *get_year_month_int(d), 1).strftime('%Y-%m-%d'))
            cur.execute(sqlstr.format(**replacements), params)
            yield cur.fetchall()
            # expense before
            sqlstr = (
                'select '
                'description as key, '
                'coalesce(sum(price * amount), 0) as value '
                'from hvhnonc_category '
                'left join ('
                'select '
                'price, '
                'hvhnonc_out.amount as amount, '
                'category '
                'from hvhnonc_out '
                'inner join hvhnonc_in '
                'on hvhnonc_out.in_ID = hvhnonc_in.ID '
                'and hvhnonc_out.unregister_date < ? '
                'and {conditions}) '
                'on description = category '
                'group by description;')
            replacements, params = get_sql_conditions(d)
            params.insert(0, datetime.date(
                *get_year_month_int(d), 1).strftime('%Y-%m-%d'))
            cur.execute(sqlstr.format(**replacements), params)
            yield cur.fetchall()
            # income in month
            sqlstr = (
                'select '
                'description as key, '
                'coalesce(sum(price * amount), 0) as value '
                'from hvhnonc_category '
                'left join hvhnonc_in '
                'on description = category '
                'and acquire_date between ? and ? '
                'and {conditions} '
                'group by description;')
            replacements, params = get_sql_conditions(d)
            magicDate = datetime.date(*get_year_month_int(d), 15)
            params.insert(0, magicDate.replace(day=1).strftime('%Y-%m-%d'))
            magicDate = magicDate.replace(day=28) + datetime.timedelta(days=4)
            magicDate -= datetime.timedelta(days=magicDate.day)
            params.insert(1, magicDate.strftime('%Y-%m-%d'))
            cur.execute(sqlstr.format(**replacements), params)
            yield cur.fetchall()
            # expenses in month
            sqlstr = (
                'select '
                'description as key, '
                'coalesce(sum(price * amount), 0) as value '
                'from hvhnonc_category '
                'left join ('
                'select '
                'price, '
                'hvhnonc_out.amount as amount, '
                'category '
                'from hvhnonc_out '
                'inner join hvhnonc_in '
                'on hvhnonc_out.in_ID = hvhnonc_in.ID '
                'and unregister_date between ? and ? '
                'and {conditions}) '
                'on description = category '
                'group by description;')
            replacements, params = get_sql_conditions(d)
            magicDate = datetime.date(*get_year_month_int(d), 15)
            params.insert(0, magicDate.replace(day=1).strftime('%Y-%m-%d'))
            magicDate = magicDate.replace(day=28) + datetime.timedelta(days=4)
            magicDate -= datetime.timedelta(days=magicDate.day)
            params.insert(1, magicDate.strftime('%Y-%m-%d'))
            cur.execute(sqlstr.format(**replacements), params)
            yield cur.fetchall()
            con.close()

        def fetch_details(d):
            """The rows of the details for monthly report.

            group by category, sort by purchase date ascended."""
            con, cur = connect._get_connection(useSQL3Row=True)
            # get all categories
            sqlstr = ('select description from hvhnonc_category;')
            cur.execute(sqlstr)
            categories = cur.fetchall()  # a list of sqlite3.Rows
            categories = [x['description'] for x in categories if x]
            # preconstruct params for sql
            conditions, tempParams = get_sql_conditions(d)
            magicMinDate = datetime.date(*get_year_month_int(d), 1)
            magicMaxDate = magicMinDate.replace(
                day=28) + datetime.timedelta(days=4)
            magicMaxDate -= datetime.timedelta(days=magicMaxDate.day)
            params = [None, magicMinDate, magicMaxDate] + list(tempParams)
            # for every category, fetch in data and out data order by pchd
            for category in categories:
                # in_data union all out_data inner join in data
                sqlstr = (
                    'select * '
                    'from ('
                    'select '
                    'category, name, brand, spec, unit, price, '
                    'amount as register_amount, '
                    '0 as unregister_amount, purchase_date, '
                    'acquire_date as date, place, remark '
                    'from hvhnonc_in '
                    'union all '
                    'select '
                    'category, name, brand, spec, unit, price, '
                    '0 as register_amount, '
                    'hvhnonc_out.amount as unregister_amount, '
                    'purchase_date, unregister_date as date, place, remark '
                    'from hvhnonc_out '
                    'inner join hvhnonc_in '
                    'on hvhnonc_out.in_ID = hvhnonc_in.ID)'
                    'where category = ?'
                    'and date between ? and ? '
                    'and {conditions} '
                    'order by purchase_date asc;')
                params[0] = category
                cur.execute(sqlstr.format(**conditions), params)
                ret = cur.fetchall()
                if ret:
                    yield category, ret
            con.close()

        def parse_data(data):
            """Parses data for excel use."""
            parsed = []
            data_p1, data_detail = data
            # income before month, expense before month,
            # income in month, expense in month
            ibm, ebm, iim, eim = data_p1
            # 1st row: title of data_p1
            parsed.append(['科目', '上月結存金額', '本月增加', '本月減少',
                           '本月結存金額'])
            # next line: the data_p1
            accumulated = ['合計', 0, 0, 0, 0]
            for r_ibm, r_ebm, r_iim, r_eim in zip(ibm, ebm, iim, eim):
                balance_before_month = r_ibm['value'] - r_ebm['value']
                income_in_month = r_iim['value']
                expense_in_month = r_eim['value']
                total_balance = (r_ibm['value'] - r_ebm['value']
                                 + r_iim['value'] - r_eim['value'])
                # put data in the parsed row
                parsed.append([r_ibm['key'], balance_before_month,
                               income_in_month, expense_in_month,
                               total_balance])
                # do the accumulation
                to_add = (balance_before_month, income_in_month,
                          expense_in_month, total_balance)
                accumulated[1:] = [sum(x)
                                   for x in zip(accumulated[1:], to_add)]
            parsed.append(accumulated)
            # insert an empty line
            parsed.append([])
            # next: title of data_details
            parsed.append(['物品名稱', '廠牌及說明', '計量單位', '單價',
                           '本月增加數量', '本月增加總價', '本月減少數量',
                           '本月減少總價', '購置日期', '存置地點', '備註事項'])
            # next line: parsed data_detail
            total = {'qty_in': 0, 'tp_in': 0, 'qty_out': 0, 'tp_out': 0}
            for category, list_grouped_by_category in data_detail:
                # insert header for category
                parsed.append([category, ])
                # accumulate quantity and total price
                accu = {'qty_in': 0, 'tp_in': 0, 'qty_out': 0, 'tp_out': 0}
                for row in list_grouped_by_category:
                    temp_row = []
                    temp_row.append(row['name'])
                    temp_row.append(row['brand'] + ' ' + row['spec'])
                    temp_row.append(row['unit'])
                    temp_row.append(row['price'])
                    if not row['register_amount']:
                        temp_row.append(0)
                        temp_row.append(0)
                    else:
                        temp_row.append(row['register_amount'])
                        temp_row.append(row['register_amount'] * row['price'])
                        accu['qty_in'] += int(row['register_amount'])
                        accu['tp_in'] += int(row['register_amount']
                                             * row['price'])
                    if not row['unregister_amount']:
                        temp_row.append(0)
                        temp_row.append(0)
                    else:
                        temp_row.append(row['unregister_amount'])
                        temp_row.append(
                            row['unregister_amount'] * row['price'])
                        accu['qty_out'] += int(row['unregister_amount'])
                        accu['tp_out'] += int(row['unregister_amount']
                                              * row['price'])
                    temp_date = row['purchase_date'].split('-')
                    temp_date = [int(x) for x in temp_date]
                    temp_date[0] -= 1911  # in ROC years
                    temp_date = [str(x).zfill(2) for x in temp_date]
                    temp_date = '-'.join(temp_date)
                    temp_row.append(temp_date)
                    temp_row.append(row['place'])
                    temp_row.append(row['remark'])
                    parsed.append(temp_row)
                # partial sum row (in column 4,5,6,7)
                parsed.append(['小計', '', '', '', ] + [accu[k]
                                                      for k in accu.keys()])
                total = {k : total[k] + accu[k] for k in total.keys()}
            # total sum
            parsed.append(['合計', '', '', '', ] + [total[k]
                                                  for k in total.keys()])
            return parsed

        def write_excel(array_2d, filename):
            """Save 2d array to .xls file."""
            wb = xlwt.Workbook()
            ws = wb.add_sheet('result')
            for rc, row in enumerate(array_2d):
                for cc, column in enumerate(row):
                    try:
                        ws.write(r=rc, c=cc, label=column)
                    except:
                        # skip if encounter problems
                        continue
            wb.save(filename)

        def split_data_for_docx(data):
            """splits data into 2, exclusively for this document."""
            for i, row in enumerate(data):
                if not row:
                    return(data[:i], data[i + 1:])

        def write_docx(array_2d, filename, d):
            """Open template and modify and save."""
            # open template
            template = Document('./mydocbuilder/monthly_report_template.docx')
            self.setMyFont(template)
            # fill in header
            header_p1 = template.sections[0].first_page_header
            header = template.sections[0].header
            # header_p1: year {y} and month {m}
            year, month = map(str, get_year_month_int(d))
            for p in header_p1.paragraphs:
                if '{' in p.text:
                    # fill in
                    p.text = p.text.format(**{'y': year, 'm': month})
            # header:
            row = header.tables[0].rows[0]
            # prepare the date for filling
            date_fill = {}
            date_fill['ys'], date_fill['ms'] = str(year), str(month)
            date_fill['ye'], date_fill['me'] = str(year), str(month)
            date_fill['ds'] = '1'
            # tricky part: the last day of month
            magic_date = datetime.date(*get_year_month_int(d), 28)
            magic_date += datetime.timedelta(days=4)
            magic_date -= datetime.timedelta(magic_date.day)
            date_fill['de'] = str(magic_date.day)
            # fill in the blanks
            for cell in row.cells:
                for p in cell.paragraphs:
                    if '{dept}' in p.text:
                        p.text = p.text.format(**{'dept': '秘書室'})
                    elif '{' in p.text:
                        p.text = p.text.format(**date_fill)
                        # this paragraph should be in center
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # split parsed data into 2
            data_p1, data_details = split_data_for_docx(data_parsed)
            # write into the tables: table_p1
            table = template.tables[0]
            for data_row in data_p1[1:]:  # skip title row
                table_row = table.add_row()
                for i, cell in enumerate(table_row.cells):
                    try:
                        cell.text = str(data_row[i])
                        if i == 0:
                            cell.paragraphs[0].alignment = \
                                    WD_PARAGRAPH_ALIGNMENT.CENTER
                        else:
                            cell.paragraphs[0].alignment = \
                                    WD_PARAGRAPH_ALIGNMENT.RIGHT
                    except IndexError:
                        pass  # ignore sparse data_row
            # write into the tables: table_details
            table = template.tables[2]
            # skip title row
            for row_count, data_row in enumerate(data_details[1:]):
                if row_count != 0:
                    table_row = table.add_row()
                else:
                    table_row = table.rows[-1]
                for i, cell in enumerate(table_row.cells):
                    try:
                        cell.text = str(data_row[i])
                        cell.paragraphs[0].alignment = \
                                WD_PARAGRAPH_ALIGNMENT.CENTER
                    except IndexError:
                        pass  # ignore sparse data_row
            # save file
            template.save(filename)

        def convert_to_pdf():
            """Convert to pdf and save (using current working directory)"""
            cwd = os.getcwd()
            self.docx_to_pdf(cwd + '\\\\result.docx', cwd + '\\\\result.pdf')

        self.status_update.emit(6, 0, 'initalizing...')  #0/6
        print('create_monthly_report')
        self.status_update.emit(6, 1, 'fetching...')  #1/6
        data_p1, data_details = fetch_from_database(self.kwargs)
        self.status_update.emit(6, 2, 'parsing...')  #2/6
        data_parsed = parse_data((data_p1, data_details))
        self.status_update.emit(6, 3, 'writing excel file...')  #3/6
        write_excel(data_parsed, 'result.xls')
        self.status_update.emit(6, 4, 'writing word file...')  #4/6
        write_docx(data_parsed, 'result.docx', self.kwargs)
        self.status_update.emit(6, 5, 'converting to pdf...')  #5/6
        convert_to_pdf()
        self.status_update.emit(6, 6, 'Done!')  #6/6

    def create_full_report(self):
        """Creates full report, saves data as excel, docx, and pdf."""

        def forge_conditions_and_params(d):
            """returns the {conditions} and according params for a table"""
            # conditions: determined by d
            replacements = {'conditions': ''}
            params = []
            for k, v in d.items():
                if 'date' in k:  # date string tuple
                    if k == 'acquire_date':  # ignore acquire date
                        pass
                    else:
                        replacements['conditions'] += \
                            '({} between ? and ?) and '.format(k)
                        params.extend(v)
                else:
                    replacements['conditions'] += \
                        ('{0} like ? and '.format(k))
                    params.append('%' + v + '%')
            replacements['conditions'] += '1'  # (cond1 and )(cond2 and )(1)
            return replacements, params

        def fetch_from_database(d):
            """Fetch data from database for full report construction.

            return:
                2 lists if sqlite3.Row: in and out.
            """
            # first fetch: all rows where conditions are met,
            #              but only type, amount and ID
            con, cur = connect._get_connection(useSQL3Row=True)
            conditions, params = forge_conditions_and_params(d)
            sqlstr = (
                    'select '
                        "'in' as type, ID, amount "
                    'from hvhnonc_in '
                    'where acquire_date between ? and ? '
                        'and {conditions} '
                    'union all '
                    'select '
                        "'out' as type, hvhnonc_in.ID as ID, "
                        'hvhnonc_out.amount as amount '
                    'from hvhnonc_out '
                    'inner join hvhnonc_in '
                        'on hvhnonc_out.in_ID = hvhnonc_in.ID '
                        'and unregister_date between ? and ? '
                        'and {conditions}')
            try:
                _, date_max = d['acquire_date']
                magic_date = (str(date_max.year), str(date_max.month).zfill(2),
                        str(date_max.day).zfill(2))
            except KeyError:
                date_temp = datetime.date.today()
                magic_date = (str(date_temp.year),
                              str(date_temp.month).zfill(2),
                              str(date_temp.day).zfill(2))
            magic_date = '-'.join(magic_date)
            params.insert(0, '1910-12-31')
            params.insert(1, magic_date)
            params.insert(2, '1910-12-31')
            params.insert(3, magic_date)
            cur.execute(sqlstr.format(**conditions), params)
            data = cur.fetchall()
            return data

        def substract_data(data):
            """The data of the difference of 'in' data and 'out' data."""
            # make a {`ID`: `amount`} dictionary for both in_data and out_data
            in_data = {row['ID']: row['amount'] for row in data if row['type'] == 'in'}
            out_data = {row['ID']: row['amount'] for row in data if row['type'] == 'out'}
            delta_data = in_data.copy()
            for id in out_data.keys():
                delta_data[id] -= out_data[id]
                if delta_data[id] <= 0:
                    delta_data.pop(id)
            return delta_data

        def fetch_detail_data(data):
            """Get the required details of the reduced id list."""
            # fetch detailed delta data
            delta_details = []
            con, cur = connect._get_connection(useSQL3Row=True)
            sqlstr = (
                    'select '
                        'ID, add_list_ID, object_ID, serial_ID, name, spec, '
                        'unit, price, acquire_date, keep_year, place, '
                        'keep_department, keeper '
                    'from hvhnonc_in '
                    'where ID = ?')
            for id in data.keys():
                cur.execute(sqlstr, (id,))
                delta_details += cur.fetchall()
            return delta_details

        def parse_data(data_from_sql, data_substracted):
            """Parse the fetched data to table, for excel printing."""
            return_data = []  # list of lists
            # title
            return_data.append(['非消耗品增加單編號', '物品編號', '物品名稱',
                                '規格', '單位', '數量', '單價', '總價',
                                '取得日期', '使用年限', '存置地點',
                                '保管或使用單位', '保管或使用人'])
            # data row
            for row in data_from_sql:
                data_row = []
                # '非消耗品增加單編號'
                data_row.append(row['add_list_ID'])
                # '物品編號'
                id = row['object_ID'] + '-' + row['serial_ID']
                id = id.replace(' ', '')
                data_row.append(id)
                # '物品名稱'
                data_row.append(row['name'])
                # '規格'
                data_row.append(row['spec'])
                # '單位'
                data_row.append(row['unit'])
                # '數量', key = ID, value = qty remain
                try:
                    data_row.append(data_substracted[row['ID']])
                except KeyError:
                    data_row.append('')
                # '單價'
                data_row.append(row['price'])
                # '總價'
                data_row.append(data_row[-1] * data_row[-2])
                # '取得日期'
                data_row.append(row['acquire_date'])
                # '使用年限'
                data_row.append(row['keep_year'])
                # '存置地點'
                data_row.append(row['place'])
                # '保管或使用單位'
                data_row.append(row['keep_department'])
                # '保管或使用人'
                data_row.append(row['keeper'])
                return_data.append(data_row)
            return return_data

        def write_to_excel(array_2d, filename):
            """Save 2d array to .xls file."""
            wb = xlwt.Workbook()
            ws = wb.add_sheet('result')
            for rc, row in enumerate(array_2d):
                for cc, column in enumerate(row):
                    try:
                        ws.write(r=rc, c=cc, label=column)
                    except:
                        # skip if encounter problems
                        continue
            wb.save(filename)

        def write_to_docx(data, filename, status_update):
            """Write data into a template and save as new."""
            template = Document('./mydocbuilder/full_report_template.docx')
            self.setMyFont(template)
            table = template.tables[0]
            for row_c, data_row in enumerate(data[1:]):
                # get a new row to write in
                if row_c == 0:  # no need for a new row
                    table_row = table.rows[-1]
                else:
                    table_row = table.add_row()
                # write in values
                for col_c, cell in enumerate(table_row.cells):
                    cell.text = str(data_row[col_c])
                status_update.emit(
                        8, 6,
                        'writing word file({}/{})...'.format(row_c,
                                                             len(data[1:])))
            template.save(filename)

        def convert_to_pdf():
            """Convert to pdf and save (using current working directory)"""
            cwd = os.getcwd()
            self.docx_to_pdf(cwd + '\\\\result.docx', cwd + '\\\\result.pdf')
            print('pdf conversion done!')

        self.status_update.emit(8, 0, 'initalizing...')  #0/8
        #print('create_full_report')
        self.status_update.emit(8, 1, 'fetching...')  #1/8
        data = fetch_from_database(self.kwargs)
        self.status_update.emit(8, 2, 'data substracting...')  #2/8
        data_substracted = substract_data(data)
        self.status_update.emit(8, 3, 'fetching detailed data...')  #3/8
        data_detail = fetch_detail_data(data_substracted)
        self.status_update.emit(8, 4, 'parsing...')  #4/8
        data_parsed = parse_data(data_detail, data_substracted)
        self.status_update.emit(8, 5, 'writing excel file...')  #5/8
        write_to_excel(data_parsed, 'result.xls')
        #status update 6/8 is done in the function
        write_to_docx(data_parsed, 'result.docx', self.status_update)
        self.status_update.emit(8, 7, 'converting to pdf...')  #7/8
        convert_to_pdf()
        self.status_update.emit(8, 8, 'Done!')  #8/8


def main():
    myDocBuilder = DocBuilder()
    myDocBuilder.construct()


if __name__ == '__main__':
    main()
